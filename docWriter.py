import pandas as pd, docx, json, re, os
from fileinput import filename
from PyQt5.QtCore import pyqtSignal, Qt, QThread
from docx.shared import RGBColor, Inches, Cm

def sheet_per_document(projectLoaded, dumpOffsets, categories_to_dump, fileDirectory):   
    default_filters = [key for key, contents in projectLoaded[list(projectLoaded.keys())[0]].items() if type(contents) == dict and key in categories_to_dump]
    writer = pd.ExcelWriter(fileDirectory, engine='xlsxwriter')


    for documentName in projectLoaded.keys():
        max_filter_length = 0
        for filterName in default_filters:
            max_filter_length= max(max_filter_length, len(projectLoaded[documentName][filterName]))


        multicolumn = []
        each_entity_in_filter_has = ["entity", "index"] if dumpOffsets == True else ["entity"]
        for filterName in default_filters:
            for miniindex in each_entity_in_filter_has:
                multicolumn.append((filterName,miniindex))
        
    
       
        multicolumn = pd.MultiIndex.from_tuples(multicolumn)


        a = pd.DataFrame(index = range(max_filter_length), columns = multicolumn)


        for filterName in default_filters:
            entityTexts = [entityText for idx, (startidx, (endidx, entityText)) in enumerate(projectLoaded[documentName][filterName].items())]
            offsets = [f"{startidx} - {endidx}" for idx, (startidx, (endidx, entityText)) in enumerate(projectLoaded[documentName][filterName].items())]

            a[(filterName, "entity")] = pd.Series(entityTexts, index = a.index[:len(entityTexts)])
            if dumpOffsets == True:
                a[(filterName, "index") ] =  pd.Series(offsets, index = a.index[:len(offsets)])
        

        #a.reset_index(inplace=True, drop=True)
        a.to_excel(writer, sheet_name=documentName[:30])

    writer.save()



def single_sheet(projectLoaded, dumpOffsets, categories_to_dump, fileDirectory):   
    default_filters = [key for key, contents in projectLoaded[list(projectLoaded.keys())[0]].items() if type(contents) == dict and key in categories_to_dump]
    multiindex = []
    each_entity_in_filter_has = ["entity", "index"] if dumpOffsets == True else ["entity"]
    for documentName in projectLoaded.keys():
        for miniindex in each_entity_in_filter_has:
            multiindex.append((documentName, miniindex))
    multiindex = pd.MultiIndex.from_tuples(multiindex, names=["documentName", "filterContents"])
    
    filters_length = {filterName : 0 for filterName in default_filters}
    for _, contentDocument in projectLoaded.items():
        for filterName in default_filters:
            filters_length[filterName] = max(filters_length[filterName], len(contentDocument[filterName]))
    
    multicolumn = []
    for filtersName, length in filters_length.items():
        for i in range(length):
            multicolumn.append((filtersName, i+1))
    multicolumn = pd.MultiIndex.from_tuples(multicolumn, names=["filterName", "#"])


    a = pd.DataFrame(index=multiindex, columns = multicolumn)


    for documentName in projectLoaded.keys():
        for filterName in default_filters:
            for idx, (startidx, (endidx, entityText)) in enumerate(projectLoaded[documentName][filterName].items()):
                a.loc[(documentName, "entity")  ,   (filterName, idx+1)] = entityText
                if dumpOffsets == True:
                    a.loc[(documentName, "index")  ,   (filterName, idx+1)] = f"{startidx} - {endidx}"
    
    a.to_excel(fileDirectory)


def excelWriterFunction(dumper, projectLoaded, directorySave, dumpOffsets, categories_to_dump, finishingfunction):
    dumper(projectLoaded, dumpOffsets, categories_to_dump, directorySave)
    finishingfunction()

def wordWriterFunction(word_function, projectLoaded, directorySave, finishingfunction):
    if not os.path.exists(directorySave):
        os.mkdir(directorySave)

    for documentName, documentContents in projectLoaded.items():
        word_function(documentName, documentContents["deidentified"], directorySave + "/" + documentName.split(".PDF")[0] + ".docx")
    
    finishingfunction()

def replace_with_category_word(filename, target, directory_save):
    red = RGBColor(0xFF, 0x0, 0x0)

    def valid_xml_char_ordinal(c):
        codepoint = ord(c)
        # conditions ordered by presumed frequency
        return (
            0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF
            )

    decode_string = lambda input_string: ''.join(c for c in input_string if valid_xml_char_ordinal(c))


    document = docx.Document()
    for section in document.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)


    document.add_heading("ANONIMIZADO_" + filename, 0)
    target = decode_string(target)
    parrafos = re.split(" - - ", target)

    for paragraphText in parrafos:
        paragraphWord = document.add_paragraph()

        pointerText = 0
        for coloredWord in re.finditer(r"[A-Z]+<+", paragraphText):
            startIdxColored, endIdxColored = coloredWord.span()
            paragraphWord.add_run(paragraphText[pointerText:startIdxColored])
            specialWordRunner = paragraphWord.add_run(paragraphText[startIdxColored:endIdxColored])
            pointerText = endIdxColored

            specialWordRunner = specialWordRunner.font
            specialWordRunner.bold = True
            specialWordRunner.underline = True
            specialWordRunner.color.rgb = red
        paragraphWord.add_run(paragraphText[pointerText:])     
   
    document.save(directory_save)

def remove_completely_word(filename, target, directory_save):
    red = RGBColor(0xFF, 0x0, 0x0)

    def valid_xml_char_ordinal(c):
        codepoint = ord(c)
        # conditions ordered by presumed frequency
        return (
            0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF
            )

    decode_string = lambda input_string: ''.join(c for c in input_string if valid_xml_char_ordinal(c))


    document = docx.Document()
    for section in document.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)


    document.add_heading("ANONIMIZADO_" + filename, 0)
    target = re.sub("[A-Z]+<+", "", target) 

  

    target = decode_string(target)
    parrafos = re.split(" - - ", target)

    for paragraphText in parrafos:
        paragraphWord = document.add_paragraph()

        pointerText = 0
        for coloredWord in re.finditer(r"[A-Z]+<+", paragraphText):
            startIdxColored, endIdxColored = coloredWord.span()
            paragraphWord.add_run(paragraphText[pointerText:startIdxColored])
            specialWordRunner = paragraphWord.add_run(paragraphText[startIdxColored:endIdxColored])
            pointerText = endIdxColored

            specialWordRunner = specialWordRunner.font
            specialWordRunner.bold = True
            specialWordRunner.underline = True
            specialWordRunner.color.rgb = red
        paragraphWord.add_run(paragraphText[pointerText:])     
   
    document.save(directory_save)

def replace_with_x_word(filename, target, directory_save):
    red = RGBColor(0xFF, 0x0, 0x0)

    def valid_xml_char_ordinal(c):
        codepoint = ord(c)
        # conditions ordered by presumed frequency
        return (
            0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF
            )

    decode_string = lambda input_string: ''.join(c for c in input_string if valid_xml_char_ordinal(c))


    document = docx.Document()
    for section in document.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)


    document.add_heading("ANONIMIZADO_" + filename, 0)

    oldlen = len(target)
    for appearnce in re.finditer("[A-Z]+<+", target) :
        startIdx, endIdx = appearnce.span()
        target = target[:startIdx] + "X"*(endIdx-startIdx-1) + "<" + target[endIdx:]

    newlen = len(target)
    if oldlen != newlen:
        print("ALERTAAAAAAAAAAAAAAAAAA")
        exit(1)

    target = decode_string(target)
    parrafos = re.split(" - - ", target)


    for paragraphText in parrafos:
        paragraphWord = document.add_paragraph()

        pointerText = 0
        for coloredWord in re.finditer(r"X+<+", paragraphText):
            startIdxColored, endIdxColored = coloredWord.span()
            paragraphWord.add_run(paragraphText[pointerText:startIdxColored])
            specialWordRunner = paragraphWord.add_run(paragraphText[startIdxColored:endIdxColored])
            pointerText = endIdxColored

            specialWordRunner = specialWordRunner.font
            specialWordRunner.bold = True
            specialWordRunner.underline = True
            specialWordRunner.color.rgb = red
        paragraphWord.add_run(paragraphText[pointerText:])     
   
    document.save(directory_save)

